import docx
import pdfplumber

def extract_text_from_document(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".docx"):
        document = docx.Document(uploaded_file)
        parts = [p.text for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n".join(parts)
    elif name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text.strip()
    return ""